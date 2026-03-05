#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для чтения DOCX файлов и извлечения их содержимого
"""

import sys
from docx import Document

def read_docx(filename):
    """Читает DOCX файл и выводит его содержимое"""
    try:
        doc = Document(filename)
        
        print(f"\n{'='*80}")
        print(f"ФАЙЛ: {filename}")
        print(f"{'='*80}\n")
        
        # Читаем параграфы
        print("ТЕКСТ ДОКУМЕНТА:")
        print("-" * 80)
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                print(f"{i}. {para.text}")
        
        # Читаем таблицы
        if doc.tables:
            print(f"\n\nТАБЛИЦЫ (найдено: {len(doc.tables)}):")
            print("-" * 80)
            
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\nТАБЛИЦА #{table_idx}")
                print(f"Размер: {len(table.rows)} строк x {len(table.columns)} колонок")
                print("-" * 40)
                
                # Выводим таблицу
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_data.append(cell_text[:30] + '...' if len(cell_text) > 30 else cell_text)
                    print(f"Строка {row_idx + 1}: {' | '.join(row_data)}")
        
        print("\n" + "="*80 + "\n")
        
    except Exception as e:
        print(f"Ошибка при чтении {filename}: {e}")
        return False
    
    return True

if __name__ == "__main__":
    files = [
        "form1.docx",
        "form2.docx",
        "form3.docx",
        "form4.docx"
    ]
    
    for filename in files:
        read_docx(filename)
