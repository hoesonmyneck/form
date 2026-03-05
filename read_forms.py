# -*- coding: utf-8 -*-
import sys
import io
from docx import Document

# Устанавливаем UTF-8 для вывода
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def analyze_form(filename):
    print(f"\n{'='*80}")
    print(f"ФАЙЛ: {filename}")
    print(f"{'='*80}\n")
    
    doc = Document(filename)
    
    # Текст
    print("ТЕКСТ:")
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"  {p.text}")
    
    # Таблицы
    print(f"\nТАБЛИЦЫ: {len(doc.tables)}")
    for i, table in enumerate(doc.tables, 1):
        print(f"\nТаблица #{i} ({len(table.rows)} строк x {len(table.columns)} колонок):")
        for j, row in enumerate(table.rows[:3], 1):  # Первые 3 строки
            cells = [cell.text.strip().replace('\n', ' ')[:40] for cell in row.cells]
            print(f"  Строка {j}: {' | '.join(cells)}")
        if len(table.rows) > 3:
            print(f"  ... и еще {len(table.rows) - 3} строк")

files = ['form1.docx', 'form2.docx', 'form3.docx', 'form4.docx']
for f in files:
    try:
        analyze_form(f)
    except Exception as e:
        print(f"Ошибка: {e}")
